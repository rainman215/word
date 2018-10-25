#! /usr/bin/env python
#coding=GB18030
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import csv 
import xlrd
from docx.enum.text import WD_ALIGN_PARAGRAPH

file = xlrd.open_workbook(u"����ע�����Ʒ��������.csv")
sh = file.sheet_by_name("Sheet1")
nrows = sh.nrows
cols=sh.ncols
# print nrows
# print sh.cell(48,1).value.encode('gbk')
document = Document()
document.add_heading('Document Title', 0)
ro=0
# heigh=500

class iput(object):
    def __init__(self,row):    
        self.hdr_cells=table.rows[row].cells
    def add_text(self,ccol,content):
        p=self.hdr_cells[ccol].paragraphs[0]
        p.add_run(content).bold=True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
#     def ex_to_word(self,col,ro,ex_col):
#         self.hdr_cells[col].text=sh.cell(ro,ex_col).value.encode('gbk') #�ַ������뵽���
    def ver_position(self,row,col):
        table.cell(row,col).vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    def ex_to_word(self,col,ro,ex_col):
        self.hdr_cells[col].text=sh.cell(ro,ex_col).value  #�������뵽���
#���ø߶�
    def set_heigh(self,totrow,heigh):
        for i in range(totrow):
            tr = table.rows[i]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '%d'%heigh)         
            trPr.append(trHeight)
        
        
while ro < 47:
    if ro > 1:
        table = document.add_table(rows=1,cols=2,style='Light List Accent 5')
        add=iput(0)
        add.add_text(0,u'�����������')
        add.ex_to_word(1,ro,3)
        add.set_heigh(0, 500)

        
        table = document.add_table(rows=3, cols=2,style='Table Grid')
        add=iput(0)
        add.add_text(0,u'��Ӧ������')
        add.ex_to_word(1, ro, 1)

        add=iput(1)
        add.add_text(0, u'������Ŀ')
        add.ex_to_word(1, ro, 4)

        add=iput(2)
        add.add_text(0,u'����Ŀ��')

        add.set_heigh(3, 500)

#         col=table.columns[0]
#         col.width=Inches(2)
#         col=table.columns[1]
#         col.width=Inches(4)
        
        
        
        table = document.add_table(rows=1, cols=1,style='Table Grid')
        hdr_cells = table.rows[0].cells
        add=iput(0)
#         add.add_text(0, u'���Թ���')
        p = hdr_cells[0].paragraphs[0]
        p.add_run(u'���Թ���').bold=True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add.ver_position(0,0)

#         hdr_cells[0].text = u'���Թ���'
        add.set_heigh(1, 500)

        table = document.add_table(rows=2, cols=4,style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'���'
        hdr_cells[1].text = u'����˵��'
        hdr_cells[2].text = u'�������'
        hdr_cells[3].text = u'ʵ����'
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text='1'
        hdr_cells[1].text=sh.cell(ro,5).value
        hdr_cells[2].text = sh.cell(ro,6).value
        add.set_heigh(2, 500)

    ro += 1
    p = document.add_paragraph()
    run = p.add_run('  ')
 
document.save('demo2.docx')

'''
# table.style = 'Light Grid Accent 3'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

col = table.columns[1]
col.width = Inches(1)

hdr_cells = table.rows[0].cells  # ��ȡ��0���������е�Ԫ��
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for i in range(3):    
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trwidth = OxmlElement('w:tcW')
    trwidth.set(qn('w:val'), "10")          
    trPr.append(trwidth)  

for i in range(3):    
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "450")          
    trPr.append(trHeight)  
    
col = table.columns[1]
col.width = Inches(5)
p = document.add_paragraph()
# p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('  ')
table = document.add_table(rows=3, cols=3,style='Table Grid')
document.save('demo.docx')
'''