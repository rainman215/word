#! /usr/bin/env python
#coding=GB18030
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
 
 
document = Document()
document.add_heading('This is my title', 0)
document.add_paragraph('my paragraph')
 
document.styles['Normal'].font.name = u'����'
p = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u'����ӵĶ������� ')
run.font.color.rgb = RGBColor(54, 95, 145)
run.font.size = Pt(36)
 
# pic = document.add_picture('logo1.PNG')
# last_paragraph = document.paragraphs[-1]
# last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ͼƬ��������
 
rows = 2
cols = 3
table = document.add_table(rows=rows, cols=cols,style = "Table Grid")  # ���2��3�еı��
 
for i in range(rows):
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "450")
    trPr.append(trHeight)  # ���߶�����
# table.autofit = False
col = table.columns[1]
col.width = Inches(5)
arr = [u'���',u"����",u"��ϸ����"]
heading_cells = table.rows[0].cells
for i in range(cols):
    p = heading_cells[i].paragraphs[0]
    run = p.add_run(arr[i])
    run.font.color.rgb = RGBColor(54, 95, 145)  # ��ɫ���ã���������RGB��ɫ
    run.font.size = Pt(12)  # �����С���ã���word������ֺ����Ӧ
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.cell(1, 1).text = u'�������'
table.add_row()
document.save('test1.docx')
